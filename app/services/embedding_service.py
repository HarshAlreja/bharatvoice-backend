"""
Embedding Service - Document Extraction, Validation & Processing Orchestrator

Embeddings are generated via HuggingFace's free hosted Inference API instead
of loading sentence-transformers/torch locally. This was changed because
loading torch locally requires more RAM than low-cost hosting tiers (e.g.
Render's free 512MB plan) provide, causing OOM (SIGKILL) crashes. Calling
a hosted API instead means this process never needs torch in memory at all.

Requires HF_TOKEN in .env -- a free HuggingFace account, Settings -> Access
Tokens -> create a "Read" token.
"""

import logging
import hashlib
import time
from typing import Tuple, Dict, Optional, Any, List
from datetime import datetime

import numpy as np
import requests
import PyPDF2
from docx import Document as DocxDocument
from flask import current_app

from app.extensions import db
from app.models.document import Document
from app.models.document_chunk import DocumentChunk
from app.vectorstore.faiss_manager import add_to_index

logger = logging.getLogger(__name__)

# ============================================================================
# CONSTANTS
# ============================================================================

MAX_FILE_SIZE_BYTES = 52428800  # 50MB
MIN_FILE_SIZE_BYTES = 1

MIN_TEXT_LENGTH_FOR_LANGUAGE_DETECTION = 200
HASH_CHUNK_SIZE = 8192

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.doc', '.xlsx', '.csv', '.json'}

DEFAULT_CHUNK_SIZE = 150     # words per chunk
DEFAULT_CHUNK_OVERLAP = 30    # words of overlap between chunks

HF_EMBEDDING_API_URL = (
    "https://router.huggingface.co/hf-inference/models/"
    "sentence-transformers/all-MiniLM-L6-v2/pipeline/feature-extraction"
)
HF_MAX_RETRIES = 5


# ============================================================================
# SERVICE
# ============================================================================

class EmbeddingService:
    """Service for document extraction, validation, chunking, and embedding."""

    EXTRACTORS = {
        '.pdf': '_extract_pdf',
        '.docx': '_extract_docx',
        '.doc': '_extract_docx',
        '.txt': '_extract_txt',
        '.xlsx': '_extract_xlsx',
        '.csv': '_extract_csv',
        '.json': '_extract_json',
    }

    CHECKERS = {
        '.pdf': '_check_pdf',
        '.docx': '_check_docx',
        '.doc': '_check_docx',
        '.txt': '_check_text',
        '.xlsx': '_check_xlsx',
        '.csv': '_check_csv',
        '.json': '_check_json',
    }

    # ========================================================================
    # CORE EXTRACTION METHODS
    # ========================================================================

    def extract_text(self, file_path: str, file_type: str) -> Optional[str]:
        """Extract text from document based on file type using dispatcher pattern."""
        file_type = self._normalize_ext(file_type)
        handler_name = self.EXTRACTORS.get(file_type)

        if not handler_name:
            logger.error(
                "Unsupported file type: %s (supported: %s)",
                file_type, ', '.join(SUPPORTED_EXTENSIONS)
            )
            return None

        handler = getattr(self, handler_name)
        return handler(file_path, file_type)

    @staticmethod
    def _normalize_ext(file_type: str) -> str:
        file_type = file_type.lower()
        return file_type if file_type.startswith('.') else f'.{file_type}'

    def _extract_pdf(self, file_path: str, file_type: str) -> Optional[str]:
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text += page.extract_text("text") + "\n"
                    except Exception as e:
                        logger.warning("Failed to extract page %d from PDF: %s", page_num, str(e))
                        continue
            return text if text.strip() else None
        except Exception as e:
            logger.error('Error extracting text from PDF (%s): %s', file_path, str(e))
            return None

    def _extract_docx(self, file_path: str, file_type: str) -> Optional[str]:
        try:
            text = ""
            doc = DocxDocument(file_path)

            for para in doc.paragraphs:
                text += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"

            return text if text.strip() else None
        except Exception as e:
            logger.error('Error extracting text from DOCX (%s): %s', file_path, str(e))
            return None

    def _extract_txt(self, file_path: str, file_type: str) -> Optional[str]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text if text.strip() else None
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return text if text.strip() else None
            except Exception as e:
                logger.error('Error extracting TXT with fallback encoding (%s): %s', file_path, str(e))
                return None
        except Exception as e:
            logger.error('Error extracting text from TXT (%s): %s', file_path, str(e))
            return None

    def _extract_xlsx(self, file_path: str, file_type: str) -> Optional[str]:
        try:
            import openpyxl
            text = ""
            workbook = openpyxl.load_workbook(file_path)

            for sheet in list(workbook.sheetnames)[:5]:
                text += f"\n--- Sheet: {sheet} ---\n"
                worksheet = workbook[sheet]

                for row_idx, row in enumerate(worksheet.iter_rows()):
                    if row_idx >= 500:
                        break
                    for cell in row:
                        if cell.value:
                            text += str(cell.value) + " | "
                    text += "\n"

            return text if text.strip() else None
        except ImportError:
            logger.warning("openpyxl not installed. Cannot extract XLSX files.")
            return None
        except Exception as e:
            logger.error('Error extracting text from XLSX (%s): %s', file_path, str(e))
            return None

    def _extract_csv(self, file_path: str, file_type: str) -> Optional[str]:
        try:
            import csv
            text = ""
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                for row_idx, row in enumerate(reader):
                    if row_idx >= 1000:
                        break
                    text += " | ".join(row) + "\n"
            return text if text.strip() else None
        except Exception as e:
            logger.error('Error extracting text from CSV (%s): %s', file_path, str(e))
            return None

    def _extract_json(self, file_path: str, file_type: str) -> Optional[str]:
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            text = json.dumps(data, indent=2, ensure_ascii=False)
            return text if text.strip() else None
        except Exception as e:
            logger.error('Error extracting text from JSON (%s): %s', file_path, str(e))
            return None

    # ========================================================================
    # CORRUPTION CHECKING
    # ========================================================================

    def check_file_corruption(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str]]:
        file_type = self._normalize_ext(file_type)
        checker_name = self.CHECKERS.get(file_type)

        if not checker_name:
            return False, f"Unsupported file type for corruption check: {file_type}"

        checker = getattr(self, checker_name)
        return checker(file_path, file_type)

    def _check_pdf(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str]]:
        try:
            with open(file_path, 'rb') as file:
                if file.read(4) != b'%PDF':
                    return False, "Invalid PDF header"
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                _ = len(reader.pages)
            return True, None
        except Exception as e:
            return False, f"PDF corruption detected: {str(e)}"

    def _check_docx(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str]]:
        try:
            import zipfile
            if not zipfile.is_zipfile(file_path):
                return False, "DOCX is not valid ZIP format"
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                if 'word/document.xml' not in zip_file.namelist():
                    return False, "Missing word/document.xml in DOCX"
            _ = DocxDocument(file_path)
            return True, None
        except Exception as e:
            return False, f"DOCX corruption detected: {str(e)}"

    def _check_text(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str]]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                _ = file.read(1024)
            return True, None
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    _ = file.read(1024)
                return True, None
            except Exception as e:
                return False, f"TXT file encoding error: {str(e)}"
        except Exception as e:
            return False, f"TXT file error: {str(e)}"

    def _check_xlsx(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str]]:
        try:
            import zipfile
            import openpyxl
            if not zipfile.is_zipfile(file_path):
                return False, "XLSX is not valid ZIP format"
            _ = openpyxl.load_workbook(file_path, read_only=True)
            return True, None
        except ImportError:
            return True, None
        except Exception as e:
            return False, f"XLSX corruption detected: {str(e)}"

    def _check_csv(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str]]:
        try:
            import csv
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                # Read up to 5 rows just to sanity-check the file parses --
                # a file with FEWER than 5 rows is still perfectly valid.
                for _, row in zip(range(5), reader):
                    pass
            return True, None
        except Exception as e:
            return False, f"CSV file error: {str(e)}"

    def _check_json(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str]]:
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as file:
                json.load(file)
            return True, None
        except json.JSONDecodeError as e:
            return False, f"JSON decode error: {str(e)}"
        except Exception as e:
            return False, f"JSON file error: {str(e)}"

    # ========================================================================
    # FILE VALIDATION
    # ========================================================================

    def validate_file(self, file_path: str, file_type: str, file_size_bytes: int) -> Tuple[bool, Optional[str]]:
        if file_size_bytes == 0:
            return False, "File is empty (0 bytes)"
        if file_size_bytes < MIN_FILE_SIZE_BYTES:
            return False, f"File is too small (minimum {MIN_FILE_SIZE_BYTES} bytes)"
        if file_size_bytes > MAX_FILE_SIZE_BYTES:
            return False, f"File exceeds maximum size ({MAX_FILE_SIZE_BYTES / 1024 / 1024:.0f}MB)"

        file_type = self._normalize_ext(file_type)
        if file_type not in SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type: {file_type}"

        return self.check_file_corruption(file_path, file_type)

    # ========================================================================
    # METADATA & HASHING
    # ========================================================================

    def calculate_file_hash(self, file_path: str) -> str:
        sha256_hash = hashlib.sha256()
        try:
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(HASH_CHUNK_SIZE), b''):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        except Exception as e:
            logger.error('Error calculating file hash: %s', str(e))
            return ""

    def detect_language(self, text: str) -> str:
        if len(text) < MIN_TEXT_LENGTH_FOR_LANGUAGE_DETECTION:
            return 'en'
        try:
            from langdetect import detect
            return detect(text)
        except Exception as e:
            logger.warning('Language detection failed: %s', str(e))
            return 'en'

    # ========================================================================
    # CHUNKING + EMBEDDING (via HuggingFace Inference API -- no local torch)
    # ========================================================================

    @staticmethod
    def _split_into_chunks(text: str, chunk_size: int = DEFAULT_CHUNK_SIZE,
                            overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
        words = text.split()
        if not words:
            return []

        chunks = []
        start = 0
        while start < len(words):
            end = start + chunk_size
            chunks.append(" ".join(words[start:end]))
            start = end - overlap if end - overlap > start else end
        return chunks

    def _generate_embedding_vector(self, text: str) -> list:
        """
        Calls HuggingFace's free Inference API instead of loading the model
        locally. Avoids the ~1GB+ RAM footprint of torch/sentence-transformers,
        which was crashing low-memory hosts (Render free tier OOM/SIGKILL).

        Requires HF_TOKEN in .env (free HuggingFace account -> Settings ->
        Access Tokens -> create a "Read" token).
        """
        headers = {"Authorization": f"Bearer {current_app.config['HF_TOKEN']}"}

        for attempt in range(HF_MAX_RETRIES):
            resp = requests.post(
                HF_EMBEDDING_API_URL,
                headers=headers,
                json={"inputs": text},
                timeout=30,
            )

            if resp.status_code == 200:
                vector = resp.json()
                arr = np.array(vector, dtype="float32")
                # Some responses come back per-token instead of pre-pooled --
                # mean-pool defensively if so.
                if arr.ndim > 1:
                    arr = arr.mean(axis=0)

                norm = np.linalg.norm(arr)
                if norm > 0:
                    arr = arr / norm
                return arr.tolist()

            if resp.status_code == 503:
                # Model is cold-starting on HF's shared infrastructure
                wait_time = 10
                try:
                    wait_time = resp.json().get("estimated_time", 10)
                except Exception:
                    pass
                logger.info("HF embedding model cold-starting, waiting %.1fs (attempt %d)", wait_time, attempt + 1)
                time.sleep(min(wait_time, 20))
                continue

            resp.raise_for_status()

        raise RuntimeError("HuggingFace Inference API did not respond after retries (model may be cold-starting)")

    # ========================================================================
    # ERROR HANDLING
    # ========================================================================

    def _error_response(self, message: str, processing_time: float,
                         file_type: str = '', file_size: int = 0, document: Optional[Document] = None) -> Dict[str, Any]:
        if document is not None:
            document.error_message = message
        return {
            'status': 'error', 'text': None, 'metadata': {}, 'hash': None,
            'file_size': file_size, 'file_type': file_type, 'chunk_count': 0,
            'error_message': message, 'processing_time': processing_time,
        }

    def _log_and_error(self, message: str, filename: str, processing_time: float,
                        file_type: str = '', file_size: int = 0, document: Optional[Document] = None) -> Dict[str, Any]:
        logger.error('Processing failed for %s (%s): %s', filename, file_type, message)
        return self._error_response(message, processing_time, file_type, file_size, document)

    # ========================================================================
    # DOCUMENT PROCESSING (full pipeline)
    # ========================================================================

    def process_document(
        self,
        business_id: int,
        document_id: int,
        file_path: str,
        file_type: str,
        original_filename: str,
        file_size_bytes: int,
    ) -> Dict[str, Any]:
        """
        Full pipeline: validate -> extract -> hash -> detect language -> chunk ->
        embed each chunk (via HF API) -> store in FAISS(business_id) -> write
        document_chunks rows.
        """
        start_time = time.time()

        try:
            document = Document.query.get(document_id)
            if not document:
                logger.error('Document %s not found', document_id)
                return self._error_response('Document not found', time.time() - start_time, file_type, file_size_bytes)

            document.status = 'processing'
            db.session.commit()

            is_valid, validation_error = self.validate_file(file_path, file_type, file_size_bytes)
            if not is_valid:
                result = self._log_and_error(
                    validation_error or 'Validation failed', original_filename,
                    time.time() - start_time, file_type, file_size_bytes, document
                )
                document.status = 'failed'
                db.session.commit()
                return result

            logger.info('Extracting text from %s (%s, %d bytes)', original_filename, file_type, file_size_bytes)
            text = self.extract_text(file_path, file_type)

            if not text or len(text.strip()) == 0:
                document.status = 'failed'
                db.session.commit()
                return self._log_and_error(
                    'No text extracted from document', original_filename,
                    time.time() - start_time, file_type, file_size_bytes, document
                )

            file_hash = self.calculate_file_hash(file_path)
            language = self.detect_language(text)

            logger.info('Chunking + embedding %s (%d chars)', original_filename, len(text))
            chunks = self._split_into_chunks(text)

            chunk_count = 0
            for idx, chunk_text in enumerate(chunks):
                try:
                    vector = self._generate_embedding_vector(chunk_text)
                except Exception as exc:
                    document.status = 'failed'
                    db.session.commit()
                    return self._log_and_error(
                        f'Embedding API error: {exc}', original_filename,
                        time.time() - start_time, file_type, file_size_bytes, document
                    )

                faiss_vector_id = add_to_index(business_id, vector)

                db.session.add(DocumentChunk(
                    document_id=document_id,
                    business_id=business_id,
                    chunk_text=chunk_text,
                    chunk_index=idx,
                    faiss_vector_id=faiss_vector_id,
                ))
                chunk_count += 1

            if chunk_count > 0:
                document.status = 'indexed'
                document.chunk_count = chunk_count
                db.session.commit()

                processing_time = time.time() - start_time
                logger.info('Document %s processed: %d chunks in %.2fs', document_id, chunk_count, processing_time)

                return {
                    'status': 'success',
                    'text': text[:500] + '...' if len(text) > 500 else text,
                    'metadata': {
                        'document_id': document_id, 'business_id': business_id,
                        'language': language, 'processed_at': datetime.utcnow().isoformat(),
                    },
                    'hash': file_hash, 'file_size': file_size_bytes, 'file_type': file_type,
                    'chunk_count': chunk_count, 'error_message': None, 'processing_time': processing_time,
                }

            document.status = 'failed'
            db.session.commit()
            return self._log_and_error(
                'Failed to create embeddings', original_filename,
                time.time() - start_time, file_type, file_size_bytes, document
            )

        except Exception as e:
            logger.error('Error processing document %s (%s): %s', document_id, original_filename, str(e))
            document = Document.query.get(document_id)
            if document:
                document.status = 'failed'
                db.session.commit()
            return self._log_and_error(str(e), original_filename, time.time() - start_time, file_type, file_size_bytes, document)

    # ========================================================================
    # DOCUMENT DELETION
    # ========================================================================

    def delete_document_embeddings(self, business_id: int, document_id: int) -> Tuple[bool, Optional[str]]:
        try:
            DocumentChunk.query.filter_by(document_id=document_id, business_id=business_id).delete()
            db.session.commit()
            logger.info('Deleted embeddings for document %s (business %s)', document_id, business_id)
            return True, None
        except Exception as e:
            db.session.rollback()
            error_msg = f'Error deleting embeddings: {str(e)}'
            logger.error('Failed to delete embeddings for document %s: %s', document_id, str(e))
            return False, error_msg


# ============================================================================
# MODULE-LEVEL FUNCTION
# ============================================================================
# rag_service.py needs to embed a customer's query string at retrieval time --
# it doesn't need the full document pipeline, just "text in, vector out"
# using the same HF API-backed model.

_embedding_service_singleton = EmbeddingService()


def generate_embedding(text: str) -> list:
    """Embed a single piece of text (e.g. a search query) using the same
    HuggingFace-hosted sentence-transformers/all-MiniLM-L6-v2 model as
    document processing."""
    return _embedding_service_singleton._generate_embedding_vector(text)